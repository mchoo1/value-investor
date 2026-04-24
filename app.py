"""
Value Investor - Flask Backend
Run: python app.py
Then open: http://localhost:5001
"""
import os
import sys
import json
import math
from datetime import datetime

# Ensure the project directory is in path
sys.path.insert(0, os.path.dirname(__file__))

# Support Railway / cloud: data dir can be overridden via DATA_DIR env var
# On Vercel the app directory is read-only — fall back to /tmp
_DATA_DIR = os.environ.get("DATA_DIR", os.path.join(os.path.dirname(__file__), "data"))
try:
    os.makedirs(_DATA_DIR, exist_ok=True)
except OSError:
    _DATA_DIR = "/tmp/valueinvestor_data"
    os.makedirs(_DATA_DIR, exist_ok=True)
os.environ.setdefault("DATA_DIR", _DATA_DIR)

from flask import Flask, jsonify, request, send_from_directory
from flask.json.provider import DefaultJSONProvider
import database as db
import stock_data as sd
import valuation as val


def _nan_to_null(obj):
    """Recursively convert float NaN / Inf to None so JSON output is always valid."""
    if isinstance(obj, dict):
        return {k: _nan_to_null(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_nan_to_null(v) for v in obj]
    if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
        return None
    return obj


class _NaNSafeJSONProvider(DefaultJSONProvider):
    """Flask 3.x JSON provider that converts NaN/Inf → null before serialisation."""
    def dumps(self, obj, **kwargs):
        return super().dumps(_nan_to_null(obj), **kwargs)


app = Flask(__name__, static_folder="static")
app.json_provider_class = _NaNSafeJSONProvider
app.json = _NaNSafeJSONProvider(app)
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = 0   # never cache static files

try:
    db.init_db()
except Exception as _e:
    print(f"[warn] db.init_db() failed: {_e}")


# ── Serve Frontend ───────────────────────────────────────────────
@app.route("/")
def index():
    return send_from_directory("static", "index.html")


@app.route("/api/health")
def health():
    """Quick ping so start.bat can wait until Flask is ready."""
    return jsonify({"status": "ok"})


@app.route("/static/<path:path>")
def static_files(path):
    return send_from_directory("static", path)


# ══════════════════════════════════════════════════════════════════
# STOCK DATA ENDPOINTS
# ══════════════════════════════════════════════════════════════════

@app.route("/api/stock/<ticker>", methods=["GET"])
def get_stock(ticker):
    """Get stock fundamentals and info."""
    info = sd.get_stock_info(ticker.upper())
    return jsonify(info)


@app.route("/api/stock/<ticker>/history", methods=["GET"])
def get_stock_history(ticker):
    """Get historical price data (monthly, 5 years)."""
    period = request.args.get("period", "5y")
    data = sd.get_price_history(ticker.upper(), period)
    return jsonify(data)


@app.route("/api/stock/<ticker>/financials", methods=["GET"])
def get_financials(ticker):
    """Get 5-year financial history."""
    data = sd.get_financial_history(ticker.upper())
    return jsonify(data)


# ══════════════════════════════════════════════════════════════════
# SCREENER
# ══════════════════════════════════════════════════════════════════

@app.route("/api/screen", methods=["POST"])
def screen():
    """
    Screen stocks based on value investing filters.
    Body: { market, filters, custom_tickers, sector }
    New strategy-aligned filters supported inside filters{}:
      min_market_cap ($B), min_fcf_yield (%), min_rev_growth (%), max_net_debt_ebitda (x)
    """
    body = request.json or {}
    market = body.get("market", "US")
    filters = body.get("filters", {})
    custom_tickers = body.get("custom_tickers", [])
    sector = body.get("sector", "All")
    if sector:
        filters["sector"] = sector

    # On Vercel (serverless) use a curated ~100-stock universe to stay within
    # the 60-second function timeout.  Full ~800-stock scan works locally.
    _on_vercel = bool(os.environ.get("VERCEL"))

    # Choose universe
    if custom_tickers:
        tickers = [t.upper() for t in custom_tickers]
    elif market == "SGX":
        tickers = sd.STI_COMPONENTS
    elif _on_vercel:
        tickers = sd.CLOUD_UNIVERSE
    else:
        tickers = sd.US_UNIVERSE

    results = sd.screen_stocks(tickers, filters)
    return jsonify({
        "results": results,
        "total": len(results),
        "screened": len(tickers),
        "cloud_mode": _on_vercel,
    })


@app.route("/api/screen/formula", methods=["GET"])
def score_formula():
    """Return the score formula weights for UI display."""
    return jsonify(sd.SCORE_FORMULA)


@app.route("/api/thesis/tickers", methods=["GET"])
def thesis_tickers():
    """Return all unique tickers that have ever had a thesis entry (prior picks)."""
    return jsonify(db.get_thesis_tickers())


# ── Research Queue ────────────────────────────────────────────────
@app.route("/api/research-queue", methods=["GET"])
def get_research_queue():
    return jsonify(db.get_research_queue())


@app.route("/api/research-queue", methods=["POST"])
def add_research_queue():
    b = request.json or {}
    db.add_to_research_queue(b.get("ticker", ""), b.get("notes", ""))
    return jsonify({"status": "ok"})


@app.route("/api/research-queue/<ticker>", methods=["DELETE"])
def remove_research_queue(ticker):
    db.remove_from_research_queue(ticker)
    return jsonify({"status": "ok"})


# ══════════════════════════════════════════════════════════════════
# INSIDER BUYING, CATALYSTS, RED FLAGS, COMPETITORS
# ══════════════════════════════════════════════════════════════════

@app.route("/api/stock/<ticker>/insiders", methods=["GET"])
def get_insiders(ticker):
    """Insider transactions (last 60 days) with net buy/sell signal."""
    return jsonify(sd.get_insider_activity(ticker.upper()))


@app.route("/api/stock/<ticker>/catalysts", methods=["GET"])
def get_catalysts(ticker):
    """Upcoming earnings, ex-div dates, and recent catalyst news."""
    return jsonify(sd.get_upcoming_catalysts(ticker.upper()))


@app.route("/api/stock/<ticker>/redflags", methods=["GET"])
def get_redflags(ticker):
    """Quantitative red flags + negative news from last 60 days."""
    return jsonify(sd.get_red_flags(ticker.upper()))


@app.route("/api/stock/<ticker>/competitors", methods=["GET"])
def get_competitors(ticker):
    """Top 5 competitors by industry + market cap, with growth metrics and recent news."""
    return jsonify(sd.get_competitors(ticker.upper()))


@app.route("/api/stock/<ticker>/moat", methods=["GET"])
def get_moat(ticker):
    """Auto-scored competitive moat rating (Wide / Narrow / None)."""
    return jsonify(sd.get_moat_rating(ticker.upper()))


@app.route("/api/stock/<ticker>/risk", methods=["GET"])
def get_risk(ticker):
    """Composite risk rating (Low / Medium / High)."""
    return jsonify(sd.get_risk_rating(ticker.upper()))


@app.route("/api/stock/<ticker>/targets", methods=["GET"])
def get_targets(ticker):
    """Bull / base / bear price targets."""
    return jsonify(sd.get_price_targets(ticker.upper()))


@app.route("/api/stock/<ticker>/metrics-history", methods=["GET"])
def get_metrics_history(ticker):
    """1Y and 5Y historical anchors: Rev Growth, Net Margin, FCF Margin, Op Margin, ROIC.
    Also returns current price, EPS TTM/Forward, shares, revenue TTM, P/E, Forward P/E."""
    return jsonify(sd.get_historical_metrics(ticker.upper()))


# ══════════════════════════════════════════════════════════════════
# VALUATION
# ══════════════════════════════════════════════════════════════════

@app.route("/api/valuation/dcf", methods=["POST"])
def dcf():
    """
    Run DCF valuation.
    Body: { current_fcf, growth_1_5, growth_6_10, wacc, terminal_growth, shares, net_debt }
    """
    b = request.json or {}
    result = val.dcf_valuation(
        current_fcf=float(b.get("current_fcf", 0)),
        growth_rate_1_5=float(b.get("growth_1_5", 10)),
        growth_rate_6_10=float(b.get("growth_6_10", 5)),
        wacc=float(b.get("wacc", 10)),
        terminal_growth=float(b.get("terminal_growth", 2.5)),
        shares_outstanding=float(b.get("shares", 1)),
        net_debt=float(b.get("net_debt", 0)),
    )
    return jsonify(result)


@app.route("/api/valuation/quick", methods=["POST"])
def quick_val():
    """Quick valuation (Graham + P/E + Lynch)."""
    b = request.json or {}
    result = val.quick_valuation(
        eps=float(b.get("eps", 0)),
        reasonable_pe=float(b.get("reasonable_pe", 15)),
        growth_rate=float(b.get("growth_rate", 0)),
    )
    return jsonify(result)


@app.route("/api/valuation/mos", methods=["POST"])
def mos():
    """Margin of safety calculation."""
    b = request.json or {}
    result = val.margin_of_safety(
        current_price=float(b.get("current_price", 0)),
        intrinsic_value=float(b.get("intrinsic_value", 0)),
    )
    return jsonify(result)


@app.route("/api/valuation/comparable", methods=["POST"])
def comparable():
    """Comparable/multiples valuation."""
    b = request.json or {}
    result = val.comparable_valuation(
        eps=float(b.get("eps", 0)),
        ebitda=float(b.get("ebitda", 0)),
        sector_pe=float(b.get("sector_pe", 15)),
        sector_ev_ebitda=float(b.get("sector_ev_ebitda", 10)),
        net_debt=float(b.get("net_debt", 0)),
        shares_outstanding=float(b.get("shares", 1)),
    )
    return jsonify(result)


# ══════════════════════════════════════════════════════════════════
# WATCHLIST
# ══════════════════════════════════════════════════════════════════

@app.route("/api/watchlist", methods=["GET"])
def get_watchlist():
    return jsonify(db.get_watchlist())


@app.route("/api/watchlist", methods=["POST"])
def add_watchlist():
    b = request.json or {}
    db.add_to_watchlist(
        ticker=b.get("ticker", ""),
        name=b.get("name", ""),
        market=b.get("market", "US"),
        notes=b.get("notes", ""),
    )
    return jsonify({"status": "ok"})


@app.route("/api/watchlist/<ticker>", methods=["DELETE"])
def remove_watchlist(ticker):
    db.remove_from_watchlist(ticker)
    return jsonify({"status": "ok"})


# ══════════════════════════════════════════════════════════════════
# PORTFOLIO
# ══════════════════════════════════════════════════════════════════

@app.route("/api/portfolio/snapshot", methods=["GET"])
def portfolio_snapshot():
    """Lightweight portfolio snapshot — raw DB data, no live price fetching."""
    positions = db.get_portfolio()
    return jsonify(positions)


@app.route("/api/portfolio", methods=["GET"])
def get_portfolio():
    positions = db.get_portfolio()
    # Index theses by ticker for O(1) join
    thesis_map = {t["ticker"]: t for t in db.get_thesis()}

    enriched = []
    for pos in positions:
        price, entry, shares = 0, pos.get("entry_price", 0) or 0, pos.get("shares", 0) or 0

        if pos.get("status") == "open":
            info = sd.get_stock_info(pos["ticker"])
            if "error" not in info:
                price = info.get("current_price", 0) or 0
                pos["current_price"]   = price
                pos["current_value"]   = round(price * shares, 2)
                pos["cost_basis"]      = round(entry * shares, 2)
                pos["gain_loss"]       = round((price - entry) * shares, 2)
                pos["gain_loss_pct"]   = round((price - entry) / entry * 100, 1) if entry else 0

        # ── Attach linked thesis ───────────────────────────────────
        thesis = thesis_map.get(pos["ticker"])
        if thesis:
            # Prefer 36m target, then thesis.target_price, then intrinsic_value
            t_price = (thesis.get("target_price_36m") or thesis.get("target_price")
                       or thesis.get("intrinsic_value") or 0)
            stop    = thesis.get("stop_loss") or 0
            cur     = price or entry

            pos["thesis"] = {
                "id":                      thesis.get("id"),
                "title":                   thesis.get("title"),
                "investment_case":         thesis.get("investment_case"),
                "risk_factors":            thesis.get("risk_factors"),
                "sell_trigger":            thesis.get("sell_trigger"),
                "key_90d_metric":          thesis.get("key_90d_metric"),
                "strategy":               thesis.get("strategy"),
                "verdict":                 thesis.get("verdict"),
                "conviction_tier":         thesis.get("conviction_tier"),
                "moat_rating":             thesis.get("moat_rating"),
                "moat_type":               thesis.get("moat_type"),
                "target_price":            t_price or None,
                "bear_target":             thesis.get("bear_target"),
                "bull_target":             thesis.get("bull_target"),
                "intrinsic_value":         thesis.get("intrinsic_value"),
                "stop_loss":               stop or None,
                "revenue_growth_assumption": thesis.get("revenue_growth_assumption"),
                "margin_assumption":       thesis.get("margin_assumption"),
                "position_size_pct":       thesis.get("position_size_pct"),
                "report_date":             thesis.get("report_date"),
                # Derived
                "upside_pct": round((t_price - cur) / cur * 100, 1) if t_price and cur else None,
                "vs_stop_pct": round((cur - stop) / stop * 100, 1) if stop and cur else None,
                "stop_breached": bool(stop and cur and cur < stop),
                "near_target":   bool(t_price and cur and cur >= t_price * 0.95),
            }
        else:
            pos["thesis"] = None

        enriched.append(pos)
    return jsonify(enriched)


@app.route("/api/portfolio", methods=["POST"])
def add_portfolio():
    b = request.json or {}
    db.add_position(
        ticker=b.get("ticker", ""),
        name=b.get("name", ""),
        entry_price=float(b.get("entry_price", 0)),
        shares=float(b.get("shares", 0)),
        entry_date=b.get("entry_date", datetime.now().strftime("%Y-%m-%d")),
        target_price=float(b.get("target_price", 0)),
        stop_loss=float(b.get("stop_loss", 0)),
        notes=b.get("notes", ""),
    )
    return jsonify({"status": "ok"})


@app.route("/api/portfolio/<int:pos_id>", methods=["PUT"])
def update_portfolio(pos_id):
    b = request.json or {}
    db.update_position(pos_id, **b)
    return jsonify({"status": "ok"})


@app.route("/api/portfolio/<int:pos_id>", methods=["DELETE"])
def delete_portfolio(pos_id):
    db.delete_position(pos_id)
    return jsonify({"status": "ok"})


# ══════════════════════════════════════════════════════════════════
# THESIS
# ══════════════════════════════════════════════════════════════════

@app.route("/api/thesis", methods=["GET"])
def get_thesis():
    ticker = request.args.get("ticker")
    return jsonify(db.get_thesis(ticker))


@app.route("/api/thesis", methods=["POST"])
def save_thesis():
    b = request.json or {}
    thesis_id = db.save_thesis(b)
    return jsonify({"status": "ok", "id": thesis_id})


@app.route("/api/thesis/<int:thesis_id>", methods=["DELETE"])
def delete_thesis(thesis_id):
    db.delete_thesis(thesis_id)
    return jsonify({"status": "ok"})


@app.route("/api/thesis/weekly-tracker", methods=["GET"])
def thesis_weekly_tracker():
    """
    Portfolio tracker — enriches active theses with live metrics, BUT only
    for tickers held in the open portfolio.  Includes thesis-breaking alerts.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed as afc

    # ── Filter to open portfolio tickers only ─────────────────────
    portfolio = db.get_portfolio()
    portfolio_tickers = {p["ticker"] for p in portfolio if p.get("status") == "open"}

    theses = db.get_thesis()
    active = [t for t in theses
              if t.get("status", "active") == "active"
              and t["ticker"] in portfolio_tickers]

    def enrich(t):
        ticker = t["ticker"]
        try:
            info      = sd.get_stock_info(ticker)
            cur_price = info.get("current_price") or 0
            entry_p   = t.get("current_price") or 0   # price at thesis creation

            # Price change since thesis written
            price_chg = round((cur_price - entry_p) / entry_p * 100, 1) if entry_p else None

            # Margin of safety vs target
            target  = t.get("target_price_36m") or t.get("target_price") or t.get("intrinsic_value") or 0
            mos_now = round((target - cur_price) / target * 100, 1) if target and cur_price else None

            # Live metrics
            cur_rev_gr  = round((info.get("revenue_growth") or 0) * 100, 1)
            cur_margin  = round((info.get("net_margin")     or 0) * 100, 1)
            cur_roe     = round((info.get("roe")            or 0) * 100, 1)
            cur_pe      = info.get("pe_ratio")

            # ── Thesis-breaking alerts ────────────────────────────
            alerts = []
            stop = t.get("stop_loss") or 0
            if stop and cur_price:
                if cur_price < stop:
                    alerts.append({"level": "danger",
                                   "msg": f"🚨 Stop loss breached — ${cur_price:.2f} < ${stop:.2f}"})
                elif cur_price < stop * 1.10:
                    alerts.append({"level": "warning",
                                   "msg": f"⚠️ Within 10% of stop loss (${stop:.2f})"})

            if target and cur_price and cur_price >= target * 0.95:
                alerts.append({"level": "success",
                                "msg": f"🎯 Within 5% of target (${target:.2f}) — consider trimming"})

            rev_assump = t.get("revenue_growth_assumption")
            if rev_assump is not None and cur_rev_gr is not None:
                if cur_rev_gr < float(rev_assump) - 5:
                    alerts.append({"level": "warning",
                                   "msg": f"📉 Rev growth {cur_rev_gr}% below thesis assumption ({rev_assump}%)"})

            margin_assump = t.get("margin_assumption")
            if margin_assump is not None and cur_margin is not None:
                if cur_margin < float(margin_assump) - 3:
                    alerts.append({"level": "warning",
                                   "msg": f"📉 Net margin {cur_margin}% below thesis assumption ({margin_assump}%)"})

            # Recent news (last 7 days)
            recent_news = []
            try:
                import yfinance as _yf, time as _time, pandas as _pd
                news_raw = _yf.Ticker(ticker).news or []
                cutoff   = _time.time() - 7 * 86400
                for n in news_raw[:10]:
                    if n.get("providerPublishTime", 0) >= cutoff:
                        recent_news.append({
                            "title": n.get("title", ""),
                            "url":   n.get("link", ""),
                            "date":  _pd.Timestamp(n["providerPublishTime"], unit="s").strftime("%Y-%m-%d"),
                        })
            except Exception:
                pass

            return {
                **t,
                "current_price_live":       cur_price,
                "price_change_since_entry": price_chg,
                "mos_now":                  mos_now,
                "target_price_effective":   target or None,
                "current_pe":               cur_pe,
                "current_roe":              cur_roe,
                "current_rev_growth":       cur_rev_gr,
                "current_net_margin":       cur_margin,
                "alerts":                   alerts,
                "recent_news":              recent_news[:5],
            }
        except Exception as e:
            return {**t, "error": str(e), "alerts": []}

    results = []
    with ThreadPoolExecutor(max_workers=8) as ex:
        futures = {ex.submit(enrich, t): t for t in active}
        for f in afc(futures):
            results.append(f.result())

    results.sort(key=lambda x: x.get("updated_date", ""), reverse=True)
    return jsonify(results)


# ══════════════════════════════════════════════════════════════════
# THESIS — DOCX AUTO-IMPORT
# ══════════════════════════════════════════════════════════════════

@app.route("/api/thesis/import-docx", methods=["POST"])
def import_thesis_from_docx():
    """Scan the workspace folder for HedgeFund*.docx files and upsert thesis rows.
    Returns a summary of what was imported / updated."""
    import glob, re

    try:
        from docx import Document
    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500

    # Look in the mounted workspace folder (Vercel: /tmp or the app directory)
    search_paths = [
        "/mnt/Stock--Stock/ValueInvestor/*.docx",
        "/mnt/Stock--Stock/ValueInvestor/**/*.docx",
        os.path.join(os.path.dirname(__file__), "*.docx"),
        "/tmp/*.docx",
    ]
    docx_files = []
    for pattern in search_paths:
        docx_files.extend(glob.glob(pattern, recursive=True))
    docx_files = list(set(docx_files))

    if not docx_files:
        return jsonify({"status": "no_files", "message": "No .docx files found in workspace"}), 200

    def _clean(text):
        return (text or "").strip()

    def _extract_val(text, keys):
        """Extract numeric value from text near any of the given key phrases."""
        for key in keys:
            pat = rf"{re.escape(key)}[:\s]+\$?([\d,\.]+)"
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                try:
                    return float(m.group(1).replace(",", ""))
                except Exception:
                    pass
        return None

    imported, updated, skipped = [], [], []

    for docx_path in sorted(docx_files):
        try:
            doc = Document(docx_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            fname = os.path.basename(docx_path)

            # Detect format — "TICKER — Name" per heading vs numbered sections
            headings = [p.text.strip() for p in doc.paragraphs
                        if p.style.name.startswith("Heading") and p.text.strip()]

            # Split into per-company blocks
            # Format A: "TICKER — Company Name" headings
            # Format B: numbered "1. Executive Summary" blocks
            blocks = []
            ticker_heading_re = re.compile(r"^([A-Z]{1,5})\s*[\u2014\-]\s+(.+)$")

            if any(ticker_heading_re.match(h) for h in headings):
                # Format A
                current_ticker, current_name, current_lines = None, None, []
                for p in doc.paragraphs:
                    m = ticker_heading_re.match(p.text.strip()) if p.style.name.startswith("Heading") else None
                    if m:
                        if current_ticker:
                            blocks.append((current_ticker, current_name, "\n".join(current_lines)))
                        current_ticker = m.group(1)
                        current_name   = m.group(2)
                        current_lines  = []
                    elif current_ticker:
                        current_lines.append(p.text)
                if current_ticker:
                    blocks.append((current_ticker, current_name, "\n".join(current_lines)))
            else:
                # Format B — extract ticker from first line of each "Executive Summary" section
                ticker_re = re.compile(r"\b([A-Z]{1,5})\b")
                sections, buf = [], []
                for p in doc.paragraphs:
                    if re.match(r"^\d+\.\s+Executive Summary", p.text.strip()):
                        if buf:
                            sections.append("\n".join(buf))
                        buf = [p.text]
                    else:
                        buf.append(p.text)
                if buf:
                    sections.append("\n".join(buf))
                for sec in sections:
                    lines = sec.splitlines()
                    # First non-empty line after heading usually has "TICKER — Name"
                    for line in lines[:5]:
                        m = ticker_heading_re.match(line.strip())
                        if m:
                            blocks.append((m.group(1), m.group(2), sec))
                            break
                    else:
                        # Fallback: scan for bold-like ticker pattern
                        all_tickers = ticker_re.findall(sec[:200])
                        if all_tickers:
                            blocks.append((all_tickers[0], "", sec))

            for ticker, company_name, text in blocks:
                ticker = ticker.upper().strip()
                if not ticker or len(ticker) > 5:
                    continue

                def _extract_section(label):
                    pat = rf"(?:{re.escape(label)})[:\s]*\n([\s\S]+?)(?=\n[A-Z][^\n]{{3,}}:|\Z)"
                    m = re.search(pat, text, re.IGNORECASE)
                    return _clean(m.group(1)) if m else ""

                investment_case = _extract_section("Investment Case") or _extract_section("Thesis") or _extract_section("Executive Summary")
                risk_factors    = _extract_section("Key Risks") or _extract_section("Risks") or _extract_section("Risk Factors")
                sell_trigger    = _extract_section("Kill Switch") or _extract_section("Sell Trigger") or _extract_section("Exit Criteria")
                catalysts       = _extract_section("Catalysts") or _extract_section("Key Catalysts")
                strategy        = _extract_section("Strategy") or _extract_section("Position Strategy")
                key_90d         = _extract_section("90-Day") or _extract_section("Near-Term") or _extract_section("Key Metric")

                # Extract numeric fields
                cur_price  = _extract_val(text, ["Current Price", "Price", "Trading at"])
                tgt_price  = _extract_val(text, ["Target Price", "Price Target", "36M Target", "Intrinsic Value"])
                stop_loss  = _extract_val(text, ["Stop Loss", "Stop-Loss", "Stop"])
                bear_tgt   = _extract_val(text, ["Bear Target", "Bear Case", "Bear Price"])
                bull_tgt   = _extract_val(text, ["Bull Target", "Bull Case", "Bull Price"])

                # Conviction tier
                conv_tier = None
                for tier in ["Tier 1", "Tier 2", "Tier 3", "High Conviction", "Medium Conviction"]:
                    if tier.lower() in text.lower():
                        conv_tier = tier; break

                # Moat
                moat_rating = None
                for mr in ["Wide", "Narrow", "None"]:
                    if re.search(rf"\b{mr}\s+Moat\b", text, re.IGNORECASE):
                        moat_rating = mr; break

                payload = {
                    "ticker":       ticker,
                    "title":        f"{ticker} — {company_name}" if company_name else ticker,
                    "investment_case": investment_case[:2000] if investment_case else None,
                    "risk_factors": risk_factors[:1500] if risk_factors else None,
                    "sell_trigger": sell_trigger[:1000] if sell_trigger else None,
                    "strategy":     strategy[:500] if strategy else None,
                    "key_90d_metric": key_90d[:500] if key_90d else None,
                    "current_price": cur_price,
                    "target_price": tgt_price,
                    "target_price_36m": tgt_price,
                    "bear_target":  bear_tgt,
                    "bull_target":  bull_tgt,
                    "stop_loss":    stop_loss,
                    "conviction_tier": conv_tier,
                    "moat_rating":  moat_rating,
                    "report_date":  fname,
                    "status":       "active",
                }
                # Remove None values so we don't overwrite good data with nulls
                payload = {k: v for k, v in payload.items() if v is not None}

                existing = db.get_thesis(ticker)
                thesis_id = db.save_thesis(payload)
                if existing:
                    updated.append(ticker)
                else:
                    imported.append(ticker)

        except Exception as exc:
            skipped.append({"file": os.path.basename(docx_path), "error": str(exc)})

    return jsonify({
        "status": "ok",
        "imported": imported,
        "updated":  updated,
        "skipped":  skipped,
        "files_scanned": [os.path.basename(f) for f in docx_files],
    })


# ══════════════════════════════════════════════════════════════════
# WEEKLY REVIEWS
# ══════════════════════════════════════════════════════════════════

@app.route("/api/reviews", methods=["GET"])
def get_reviews():
    ticker = request.args.get("ticker")
    thesis_id = request.args.get("thesis_id")
    return jsonify(db.get_weekly_reviews(ticker, int(thesis_id) if thesis_id else None))


@app.route("/api/reviews", methods=["POST"])
def save_review():
    b = request.json or {}
    db.save_weekly_review(b)
    return jsonify({"status": "ok"})


# ══════════════════════════════════════════════════════════════════
# MARKET OVERVIEW (for dashboard) — fast batch fetch
# ══════════════════════════════════════════════════════════════════

@app.route("/api/market/overview", methods=["GET"])
def market_overview():
    """Get major index prices via single batch yf.download call (fast)."""
    result = sd.get_index_prices()
    return jsonify(result)


@app.route("/api/market/buffett-indicator", methods=["GET"])
def buffett_indicator():
    """Buffett Indicator: US market cap / GDP, sourced from FRED public CSV endpoints.
    Cached for 6 hours — GDP is quarterly, market cap updates daily."""
    import time as _time
    from urllib.request import urlopen

    CACHE_KEY = "__buffett_indicator__"
    CACHE_TTL  = 6 * 3600  # 6 hours

    with sd._cache_lock:
        entry = sd._cache.get(CACHE_KEY)
    if entry and (_time.time() - entry["ts"]) < CACHE_TTL:
        return jsonify(entry["data"])

    try:
        def fetch_fred_csv(series_id):
            url = f"https://fred.stlouisfed.org/graph/fredgraph.csv?id={series_id}"
            with urlopen(url, timeout=12) as resp:
                text = resp.read().decode("utf-8")
            rows = []
            for line in text.strip().splitlines()[1:]:
                parts = line.strip().split(",")
                if len(parts) == 2 and parts[1] != ".":
                    rows.append({"date": parts[0], "value": float(parts[1])})
            return rows

        mc_data  = fetch_fred_csv("WILL5000INDFC")  # billions USD, daily
        gdp_data = fetch_fred_csv("GDP")             # billions USD, quarterly, annualised

        if not mc_data or not gdp_data:
            return jsonify({"error": "FRED returned no data"}), 502

        latest_mc  = mc_data[-1]
        latest_gdp = gdp_data[-1]
        ratio      = round(latest_mc["value"] / latest_gdp["value"] * 100, 1)

        # 1-year-ago ratio
        from datetime import datetime as _dt, timedelta as _td
        one_yr_ago = (_dt.now() - _td(days=365)).strftime("%Y-%m-%d")
        mc_1y  = next((d for d in reversed(mc_data)  if d["date"] <= one_yr_ago), mc_data[0])
        gdp_1y = next((d for d in reversed(gdp_data) if d["date"] <= one_yr_ago), gdp_data[0])
        ratio_1y = round(mc_1y["value"] / gdp_1y["value"] * 100, 1)

        # Historical series (quarterly, last 32 quarters ≈ 8 years) for sparkline
        history = []
        for gdp_pt in gdp_data[-32:]:
            mc_pt = next((d for d in reversed(mc_data) if d["date"] <= gdp_pt["date"]), None)
            if mc_pt:
                history.append({"date": gdp_pt["date"],
                                 "ratio": round(mc_pt["value"] / gdp_pt["value"] * 100, 1)})

        # Zone classification
        if ratio < 75:
            zone, zone_color = "Undervalued",              "emerald"
        elif ratio < 100:
            zone, zone_color = "Fair Value",               "yellow"
        elif ratio < 130:
            zone, zone_color = "Overvalued",               "orange"
        elif ratio < 175:
            zone, zone_color = "Significantly Overvalued", "red"
        else:
            zone, zone_color = "Strongly Overvalued",      "red"

        # Interpretation blurb
        blurbs = {
            "Undervalued":              "Market appears cheap relative to the economy — historically a good entry window for long-term investors.",
            "Fair Value":               "Market is broadly in line with economic output. Stock-picking matters more than macro timing here.",
            "Overvalued":               "Market is running ahead of GDP. Expect lower future returns; a margin-of-safety approach is prudent.",
            "Significantly Overvalued": "Valuations are stretched. Buffett has historically held cash or been cautious at these levels.",
            "Strongly Overvalued":      "Extreme overvaluation — above the dot-com bubble peak. Risk management is paramount.",
        }

        data = {
            "ratio":          ratio,
            "ratio_1y":       ratio_1y,
            "change_1y":      round(ratio - ratio_1y, 1),
            "zone":           zone,
            "zone_color":     zone_color,
            "market_cap_b":   round(latest_mc["value"],  0),
            "gdp_b":          round(latest_gdp["value"], 0),
            "market_cap_date": latest_mc["date"],
            "gdp_date":        latest_gdp["date"],
            "history":        history,
            "blurb":          blurbs[zone],
        }

        with sd._cache_lock:
            sd._cache[CACHE_KEY] = {"ts": _time.time(), "data": data}

        return jsonify(data)

    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/cache/clear", methods=["POST"])
def clear_cache():
    """Clear the in-memory data cache (forces fresh fetch next request)."""
    import stock_data
    with stock_data._cache_lock:
        stock_data._cache.clear()
    return jsonify({"status": "ok", "message": "Cache cleared"})


def _open_browser():
    """Wait until Flask is ready, then open the browser automatically."""
    import threading, time, webbrowser
    try:
        from urllib.request import urlopen
    except ImportError:
        import urllib2 as urlopen  # Python 2 fallback (unlikely)

    def _wait_and_open():
        for _ in range(30):          # try for up to 60 seconds
            time.sleep(2)
            try:
                urlopen("http://localhost:5001/api/health", timeout=2)
                webbrowser.open("http://localhost:5001")
                return
            except Exception:
                pass

    t = threading.Thread(target=_wait_and_open, daemon=True)
    t.start()


if __name__ == "__main__":
    # When running locally, open browser automatically
    # When running on Railway/cloud (PORT env set), skip browser open
    port = int(os.environ.get("PORT", 5001))
    is_cloud = "PORT" in os.environ

    print("\n" + "="*60)
    print("  Value Investor App")
    if is_cloud:
        print(f"  Running on port {port} (cloud mode)")
    else:
        print("  Starting... browser will open automatically.")
    print("  Press Ctrl+C to stop.")
    print("="*60 + "\n")

    if not is_cloud:
        _open_browser()

    app.run(debug=False, port=port, host="0.0.0.0", threaded=True)
